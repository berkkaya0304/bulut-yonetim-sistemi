import streamlit as st
import os
from dotenv import load_dotenv
import subprocess
import json
from datetime import datetime

# Sayfa yapÄ±landÄ±rmasÄ±
st.set_page_config(
    page_title="Huawei Cloud YÃ¶netim Paneli",
    page_icon="â˜ï¸",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Stilleri
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #FF4B4B;
        color: white;
    }
    .stButton>button:hover {
        background-color: #FF6B6B;
    }
    .css-1d391kg {
        padding: 1rem;
    }
    .stSelectbox {
        padding: 0.5rem;
    }
    .stTextInput>div>div>input {
        border-radius: 5px;
    }
    .stTextArea>div>div>textarea {
        border-radius: 5px;
    }
    .success-box {
        background-color: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        color: #721c24;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #cce5ff;
        color: #004085;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .metric-card {
        background-color: white;
        padding: 1rem;
        border-radius: 5px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# BaÅŸlÄ±k ve Ãœst Bilgi
st.title("â˜ï¸ Huawei Cloud YÃ¶netim Paneli")

# Yan menÃ¼
with st.sidebar:
    # Kimlik Bilgileri Butonu
    if st.button("ğŸ”‘ Kimlik Bilgilerini YÃ¶net", use_container_width=True):
        with st.form("credentials_form"):
            st.markdown("### ğŸ”‘ Kimlik Bilgileri")
            col1, col2 = st.columns(2)
            with col1:
                access_key = st.text_input("Access Key", type="password")
                secret_key = st.text_input("Secret Key", type="password")
            with col2:
                region = st.selectbox("BÃ¶lge", ["eu-west-101", "ap-southeast-1", "na-mexico-1"])
                project_id = st.text_input("Proje ID")
            
            submitted = st.form_submit_button("Kimlik Bilgilerini GÃ¼ncelle")
            if submitted:
                # Kimlik bilgilerini .env dosyasÄ±na kaydet
                with open(".env", "w") as f:
                    f.write(f"HUAWEI_ACCESS_KEY={access_key}\n")
                    f.write(f"HUAWEI_SECRET_KEY={secret_key}\n")
                    f.write(f"HUAWEI_REGION={region}\n")
                    f.write(f"HUAWEI_PROJECT_ID={project_id}\n")
                st.success("Kimlik bilgileri gÃ¼ncellendi!")
    
    st.markdown("---")
    
    # Ana MenÃ¼ Kategorileri
    st.markdown("### ğŸ“‹ Ana MenÃ¼")
    
    # Kategori seÃ§imi
    category = st.selectbox(
        "Kategori SeÃ§in",
        [
            "Hesaplama ve AÄŸ",
            "Depolama ve VeritabanÄ±",
            "GÃ¼venlik ve Kimlik",
            "Konteyner ve Orkestrasyon",
            "AI ve BÃ¼yÃ¼k Veri",
            "GeliÅŸtirici AraÃ§larÄ±",
            "Ä°zleme ve YÃ¶netim"
        ]
    )
    
    # Servis seÃ§imi
    if category == "Hesaplama ve AÄŸ":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "ECS YÃ¶netimi",
                "VPC YÃ¶netimi",
                "ELB YÃ¶netimi",
                "EIP YÃ¶netimi",
                "NAT Gateway",
                "VPN Gateway",
                "Direct Connect",
                "Route Table"
            ]
        )
    elif category == "Depolama ve VeritabanÄ±":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "OBS YÃ¶netimi",
                "SFS YÃ¶netimi",
                "RDS YÃ¶netimi",
                "DCS YÃ¶netimi",
                "DMS YÃ¶netimi",
                "GaussDB",
                "DWS YÃ¶netimi",
                "CSS YÃ¶netimi"
            ]
        )
    elif category == "GÃ¼venlik ve Kimlik":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "GÃ¼venlik Grubu YÃ¶netimi",
                "WAF YÃ¶netimi",
                "Anti-DDoS",
                "SSL SertifikalarÄ±",
                "IAM YÃ¶netimi",
                "KMS YÃ¶netimi",
                "CBR YÃ¶netimi"
            ]
        )
    elif category == "Konteyner ve Orkestrasyon":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "CCE YÃ¶netimi",
                "SWR YÃ¶netimi",
                "ServiceStage",
                "FunctionGraph"
            ]
        )
    elif category == "AI ve BÃ¼yÃ¼k Veri":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "ModelArts",
                "MRS YÃ¶netimi",
                "DLS YÃ¶netimi",
                "DGC YÃ¶netimi",
                "CloudTable",
                "DIS YÃ¶netimi"
            ]
        )
    elif category == "GeliÅŸtirici AraÃ§larÄ±":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "DevCloud",
                "CodeArts",
                "API Gateway",
                "ServiceComb",
                "CloudIDE"
            ]
        )
    elif category == "Ä°zleme ve YÃ¶netim":
        service = st.selectbox(
            "Servis SeÃ§in",
            [
                "Cloud Eye",
                "AOM YÃ¶netimi",
                "LTS YÃ¶netimi",
                "CES YÃ¶netimi",
                "Terraform Durumu"
            ]
        )

    # HÄ±zlÄ± Ä°ÅŸlemler
    st.markdown("### âš¡ HÄ±zlÄ± Ä°ÅŸlemler")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("ğŸ”„ Yenile"):
            subprocess.run(["terraform", "refresh"])
            st.success("Durum yenilendi!")
    with col2:
        if st.button("ğŸ“Š Rapor"):
            try:
                from docx import Document
                from docx.shared import Inches
                import os
                from datetime import datetime
                
                # Word dosyasÄ± oluÅŸtur
                doc = Document()
                
                # BaÅŸlÄ±k
                doc.add_heading('Huawei Cloud Ä°ÅŸlem Raporu', 0)
                doc.add_paragraph(f'OluÅŸturulma Tarihi: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                
                # Terraform durumunu al
                try:
                    # Ã–nce terraform init Ã§alÄ±ÅŸtÄ±r
                    init_result = subprocess.run(["terraform", "init"], capture_output=True, text=True)
                    if init_result.returncode != 0:
                        doc.add_heading('Terraform BaÅŸlatma HatasÄ±', level=2)
                        doc.add_paragraph(init_result.stderr)
                    else:
                        # Terraform durumunu al
                        state_result = subprocess.run(["terraform", "state", "list"], capture_output=True, text=True)
                        if state_result.returncode == 0:
                            doc.add_heading('Terraform Durumu', level=2)
                            doc.add_paragraph(state_result.stdout)
                        else:
                            doc.add_heading('Terraform Durumu HatasÄ±', level=2)
                            doc.add_paragraph(state_result.stderr)
                except Exception as terraform_error:
                    doc.add_heading('Terraform HatasÄ±', level=2)
                    doc.add_paragraph(str(terraform_error))
                
                # Son iÅŸlemler bÃ¶lÃ¼mÃ¼
                doc.add_heading('Son Ä°ÅŸlemler', level=1)
                doc.add_paragraph(f'Son gÃ¼ncelleme: {datetime.now().strftime("%H:%M:%S")}')
                
                # DosyayÄ± kaydet
                report_name = f"huawei_cloud_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                doc.save(report_name)
                
                st.success(f"Rapor oluÅŸturuldu: {report_name}")
                
                # DosyayÄ± indirme linki oluÅŸtur
                with open(report_name, "rb") as file:
                    st.download_button(
                        label="ğŸ“¥ Raporu Ä°ndir",
                        data=file,
                        file_name=report_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
            except Exception as e:
                st.error(f"Rapor oluÅŸturulurken hata oluÅŸtu: {str(e)}")
    
    # Durum Bilgisi
    st.markdown("---")
    st.markdown("### ğŸ“Š Sistem Durumu")
    load_dotenv()
    if os.getenv("HUAWEI_ACCESS_KEY") and os.getenv("HUAWEI_SECRET_KEY"):
        st.success("âœ… Kimlik bilgileri aktif")
    else:
        st.error("âŒ Kimlik bilgileri eksik")
    
    # Son Ä°ÅŸlemler
    st.markdown("### ğŸ“ Son Ä°ÅŸlemler")
    st.markdown(f"Son gÃ¼ncelleme: {datetime.now().strftime('%H:%M:%S')}")

# Hesaplama ve AÄŸ Kategorisi
if service == "EIP YÃ¶netimi":
    st.header("ğŸŒ EIP YÃ¶netimi")
    
    eip_action = st.selectbox(
        "EIP Ä°ÅŸlemi SeÃ§in",
        ["EIP OluÅŸtur", "EIP Listele", "EIP Sil", "EIP BaÄŸla", "EIP AyÄ±r"]
    )
    
    if eip_action == "EIP OluÅŸtur":
        with st.form("eip_form"):
            col1, col2 = st.columns(2)
            with col1:
                eip_name = st.text_input("EIP AdÄ±")
                bandwidth = st.number_input("Bant GeniÅŸliÄŸi (Mbps)", min_value=1, value=100)
            with col2:
                eip_type = st.selectbox("EIP Tipi", ["5_bgp", "5_sbgp"])
                charge_mode = st.selectbox("Ã–deme Modu", ["traffic", "bandwidth"])
            
            submitted = st.form_submit_button("EIP OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="eip_name={eip_name}" -var="bandwidth={bandwidth}"
                -var="eip_type={eip_type}" -var="charge_mode={charge_mode}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("EIP oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "NAT Gateway":
    st.header("ğŸŒ NAT Gateway YÃ¶netimi")
    
    nat_action = st.selectbox(
        "NAT Gateway Ä°ÅŸlemi SeÃ§in",
        ["NAT Gateway OluÅŸtur", "NAT Gateway Listele", "NAT Gateway Sil", "SNAT KuralÄ± Ekle", "DNAT KuralÄ± Ekle"]
    )
    
    if nat_action == "NAT Gateway OluÅŸtur":
        with st.form("nat_form"):
            col1, col2 = st.columns(2)
            with col1:
                nat_name = st.text_input("NAT Gateway AdÄ±")
                nat_type = st.selectbox("NAT Gateway Tipi", ["1", "2", "3", "4"])
            with col2:
                vpc_id = st.text_input("VPC ID")
                subnet_id = st.text_input("Subnet ID")
            
            submitted = st.form_submit_button("NAT Gateway OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="nat_name={nat_name}" -var="nat_type={nat_type}"
                -var="vpc_id={vpc_id}" -var="subnet_id={subnet_id}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("NAT Gateway oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "VPN Gateway":
    st.header("ğŸ”’ VPN Gateway YÃ¶netimi")
    
    vpn_action = st.selectbox(
        "VPN Gateway Ä°ÅŸlemi SeÃ§in",
        ["VPN Gateway OluÅŸtur", "VPN Gateway Listele", "VPN Gateway Sil", "VPN BaÄŸlantÄ±sÄ± Ekle"]
    )
    
    if vpn_action == "VPN Gateway OluÅŸtur":
        with st.form("vpn_form"):
            col1, col2 = st.columns(2)
            with col1:
                vpn_name = st.text_input("VPN Gateway AdÄ±")
                vpn_type = st.selectbox("VPN Gateway Tipi", ["Basic", "Professional"])
            with col2:
                vpc_id = st.text_input("VPC ID")
                subnet_id = st.text_input("Subnet ID")
            
            submitted = st.form_submit_button("VPN Gateway OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="vpn_name={vpn_name}" -var="vpn_type={vpn_type}"
                -var="vpc_id={vpc_id}" -var="subnet_id={subnet_id}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("VPN Gateway oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Hesaplama ve AÄŸ Kategorisi - Eksik Servisler
elif service == "ECS YÃ¶netimi":
    st.header("ğŸ’» ECS YÃ¶netimi")
    
    ecs_action = st.selectbox(
        "ECS Ä°ÅŸlemi SeÃ§in",
        ["Instance OluÅŸtur", "Instance Listele", "Instance Sil", "Instance Yeniden BaÅŸlat", "Instance Durdur"]
    )
    
    if ecs_action == "Instance OluÅŸtur":
        with st.form("ecs_form"):
            col1, col2 = st.columns(2)
            with col1:
                instance_name = st.text_input("Instance AdÄ±")
                image_id = st.text_input("Image ID")
                flavor = st.selectbox("Instance Tipi", ["s3.large.2", "s3.xlarge.2", "s3.2xlarge.2"])
            with col2:
                key_pair = st.text_input("SSH Key Pair")
                security_group = st.text_input("GÃ¼venlik Grubu")
                availability_zone = st.selectbox("BÃ¶lge", ["eu-west-101a", "eu-west-101b"])
            
            submitted = st.form_submit_button("Instance OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="instance_name={instance_name}" -var="image_id={image_id}"
                -var="flavor={flavor}" -var="key_pair={key_pair}" -var="security_group={security_group}"
                -var="availability_zone={availability_zone}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("ECS instance oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "VPC YÃ¶netimi":
    st.header("ğŸŒ VPC YÃ¶netimi")
    
    vpc_action = st.selectbox(
        "VPC Ä°ÅŸlemi SeÃ§in",
        ["VPC OluÅŸtur", "VPC Listele", "VPC Sil", "Subnet Ekle"]
    )
    
    if vpc_action == "VPC OluÅŸtur":
        with st.form("vpc_form"):
            col1, col2 = st.columns(2)
            with col1:
                vpc_name = st.text_input("VPC AdÄ±")
                cidr = st.text_input("CIDR", value="192.168.0.0/16")
            with col2:
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("VPC OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="vpc_name={vpc_name}" -var="cidr={cidr}"
                -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("VPC oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "Direct Connect":
    st.header("ğŸ”Œ Direct Connect YÃ¶netimi")
    
    directconnect_action = st.selectbox(
        "Direct Connect Ä°ÅŸlemi SeÃ§in",
        ["BaÄŸlantÄ± OluÅŸtur", "BaÄŸlantÄ± Listele", "BaÄŸlantÄ± Sil", "Virtual Interface Ekle"]
    )
    
    if directconnect_action == "BaÄŸlantÄ± OluÅŸtur":
        with st.form("directconnect_form"):
            col1, col2 = st.columns(2)
            with col1:
                connection_name = st.text_input("BaÄŸlantÄ± AdÄ±")
                bandwidth = st.selectbox("Bant GeniÅŸliÄŸi", ["1Gbps", "10Gbps", "40Gbps"])
            with col2:
                location = st.text_input("Lokasyon")
                provider = st.text_input("Servis SaÄŸlayÄ±cÄ±")
            
            submitted = st.form_submit_button("BaÄŸlantÄ± OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="connection_name={connection_name}" -var="bandwidth={bandwidth}"
                -var="location={location}" -var="provider={provider}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("Direct Connect baÄŸlantÄ±sÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "Route Table":
    st.header("ğŸ›£ï¸ Route Table YÃ¶netimi")
    
    routetable_action = st.selectbox(
        "Route Table Ä°ÅŸlemi SeÃ§in",
        ["Route Table OluÅŸtur", "Route Table Listele", "Route Table Sil", "Route Ekle"]
    )
    
    if routetable_action == "Route Table OluÅŸtur":
        with st.form("routetable_form"):
            col1, col2 = st.columns(2)
            with col1:
                table_name = st.text_input("Route Table AdÄ±")
                vpc_id = st.text_input("VPC ID")
            with col2:
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Route Table OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="table_name={table_name}" -var="vpc_id={vpc_id}"
                -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("Route Table oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Depolama ve VeritabanÄ± Kategorisi
elif service == "OBS YÃ¶netimi":
    st.header("ğŸ“¦ Object Storage Service YÃ¶netimi")
    
    obs_action = st.selectbox(
        "OBS Ä°ÅŸlemi SeÃ§in",
        ["Bucket OluÅŸtur", "Bucket Listele", "Bucket Sil", "Object YÃ¼kle"]
    )
    
    if obs_action == "Bucket OluÅŸtur":
        with st.form("obs_form"):
            col1, col2 = st.columns(2)
            with col1:
                bucket_name = st.text_input("Bucket AdÄ±")
                storage_class = st.selectbox("Depolama SÄ±nÄ±fÄ±", ["Standard", "Infrequent Access", "Archive"])
            with col2:
                region = st.selectbox("BÃ¶lge", ["eu-west-101", "ap-southeast-1", "na-mexico-1"])
                versioning = st.checkbox("Versiyonlama Aktif", value=False)
            
            submitted = st.form_submit_button("Bucket OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="bucket_name={bucket_name}" -var="storage_class={storage_class}"
                -var="region={region}" -var="versioning={versioning}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("OBS bucket oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "SFS YÃ¶netimi":
    st.header("ğŸ“ Scalable File Service YÃ¶netimi")
    
    sfs_action = st.selectbox(
        "SFS Ä°ÅŸlemi SeÃ§in",
        ["Dosya Sistemi OluÅŸtur", "Dosya Sistemi Listele", "Dosya Sistemi Sil", "EriÅŸim KuralÄ± Ekle"]
    )
    
    if sfs_action == "Dosya Sistemi OluÅŸtur":
        with st.form("sfs_form"):
            col1, col2 = st.columns(2)
            with col1:
                fs_name = st.text_input("Dosya Sistemi AdÄ±")
                size = st.number_input("Boyut (GB)", min_value=10, value=100)
            with col2:
                protocol = st.selectbox("Protokol", ["NFS", "CIFS"])
                vpc_id = st.text_input("VPC ID")
            
            submitted = st.form_submit_button("Dosya Sistemi OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="fs_name={fs_name}" -var="size={size}"
                -var="protocol={protocol}" -var="vpc_id={vpc_id}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("SFS dosya sistemi oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "RDS YÃ¶netimi":
    st.header("ğŸ—„ï¸ Relational Database Service YÃ¶netimi")
    
    rds_action = st.selectbox(
        "RDS Ä°ÅŸlemi SeÃ§in",
        ["Instance OluÅŸtur", "Instance Listele", "Instance Sil", "Yedekleme OluÅŸtur"]
    )
    
    if rds_action == "Instance OluÅŸtur":
        with st.form("rds_form"):
            col1, col2 = st.columns(2)
            with col1:
                instance_name = st.text_input("Instance AdÄ±")
                engine = st.selectbox("VeritabanÄ± Motoru", ["MySQL", "PostgreSQL", "SQL Server"])
            with col2:
                version = st.selectbox("Versiyon", ["5.7", "8.0", "12.0"])
                flavor = st.selectbox("Instance Tipi", ["rds.mysql.s1.large", "rds.mysql.s1.xlarge"])
            
            submitted = st.form_submit_button("Instance OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="instance_name={instance_name}" -var="engine={engine}"
                -var="version={version}" -var="flavor={flavor}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("RDS instance oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "GaussDB":
    st.header("ğŸ—„ï¸ GaussDB YÃ¶netimi")
    
    gauss_action = st.selectbox(
        "GaussDB Ä°ÅŸlemi SeÃ§in",
        ["Instance OluÅŸtur", "Instance Listele", "Instance Sil", "Yedekleme OluÅŸtur"]
    )
    
    if gauss_action == "Instance OluÅŸtur":
        with st.form("gauss_form"):
            col1, col2 = st.columns(2)
            with col1:
                instance_name = st.text_input("Instance AdÄ±")
                db_type = st.selectbox("VeritabanÄ± Tipi", ["MySQL", "PostgreSQL"])
            with col2:
                version = st.selectbox("Versiyon", ["8.0", "9.2"])
                flavor = st.selectbox("Instance Tipi", ["gaussdb.mysql.s1.large", "gaussdb.mysql.s1.xlarge"])
            
            submitted = st.form_submit_button("Instance OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="instance_name={instance_name}" -var="db_type={db_type}"
                -var="version={version}" -var="flavor={flavor}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("GaussDB instance oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "DWS YÃ¶netimi":
    st.header("ğŸ“Š Data Warehouse Service YÃ¶netimi")
    
    dws_action = st.selectbox(
        "DWS Ä°ÅŸlemi SeÃ§in",
        ["Cluster OluÅŸtur", "Cluster Listele", "Cluster Sil", "Yedekleme OluÅŸtur"]
    )
    
    if dws_action == "Cluster OluÅŸtur":
        with st.form("dws_form"):
            col1, col2 = st.columns(2)
            with col1:
                cluster_name = st.text_input("Cluster AdÄ±")
                node_type = st.selectbox("Node Tipi", ["dws.m3.xlarge", "dws.m3.2xlarge"])
            with col2:
                node_count = st.number_input("Node SayÄ±sÄ±", min_value=3, value=3)
                version = st.selectbox("Versiyon", ["8.1.3", "8.2.0"])
            
            submitted = st.form_submit_button("Cluster OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="cluster_name={cluster_name}" -var="node_type={node_type}"
                -var="node_count={node_count}" -var="version={version}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DWS cluster oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "CSS YÃ¶netimi":
    st.header("ğŸ” Cloud Search Service YÃ¶netimi")
    
    css_action = st.selectbox(
        "CSS Ä°ÅŸlemi SeÃ§in",
        ["Cluster OluÅŸtur", "Cluster Listele", "Cluster Sil", "Index OluÅŸtur"]
    )
    
    if css_action == "Cluster OluÅŸtur":
        with st.form("css_form"):
            col1, col2 = st.columns(2)
            with col1:
                cluster_name = st.text_input("Cluster AdÄ±")
                node_type = st.selectbox("Node Tipi", ["ess.spec-2u8g", "ess.spec-4u16g"])
            with col2:
                node_count = st.number_input("Node SayÄ±sÄ±", min_value=1, value=3)
                version = st.selectbox("Versiyon", ["7.6.2", "7.9.3"])
            
            submitted = st.form_submit_button("Cluster OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="cluster_name={cluster_name}" -var="node_type={node_type}"
                -var="node_count={node_count}" -var="version={version}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("CSS cluster oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# GÃ¼venlik ve Kimlik Kategorisi
elif service == "GÃ¼venlik Grubu YÃ¶netimi":
    st.header("ğŸ›¡ï¸ GÃ¼venlik Grubu YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["GÃ¼venlik Grubu OluÅŸtur", "GÃ¼venlik GruplarÄ±nÄ± Listele", "GÃ¼venlik Grubu Sil", "Kural Ekle"]
    )
    
    if operation == "GÃ¼venlik Grubu OluÅŸtur":
        with st.form("security_group_create"):
            name = st.text_input("GÃ¼venlik Grubu AdÄ±")
            description = st.text_area("AÃ§Ä±klama")
            vpc_id = st.text_input("VPC ID")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="security_group_name={name}" -var="description={description}" -var="vpc_id={vpc_id}"
                """, language="bash")
                st.success("GÃ¼venlik grubu oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "GÃ¼venlik GruplarÄ±nÄ± Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_networking_secgroup
            """, language="bash")
            st.success("GÃ¼venlik gruplarÄ± listeleniyor!")
            
    elif operation == "GÃ¼venlik Grubu Sil":
        with st.form("security_group_delete"):
            group_id = st.text_input("GÃ¼venlik Grubu ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_networking_secgroup.{group_id}
                """, language="bash")
                st.success("GÃ¼venlik grubu silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Kural Ekle":
        with st.form("security_group_rule_add"):
            group_id = st.text_input("GÃ¼venlik Grubu ID")
            direction = st.selectbox("YÃ¶n", ["ingress", "egress"])
            protocol = st.selectbox("Protokol", ["tcp", "udp", "icmp"])
            port_range_min = st.number_input("BaÅŸlangÄ±Ã§ Portu", min_value=1, max_value=65535)
            port_range_max = st.number_input("BitiÅŸ Portu", min_value=1, max_value=65535)
            remote_ip_prefix = st.text_input("Uzak IP Prefix")
            
            if st.form_submit_button("Kural Ekle"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="security_group_id={group_id}" -var="direction={direction}" -var="protocol={protocol}" -var="port_range_min={port_range_min}" -var="port_range_max={port_range_max}" -var="remote_ip_prefix={remote_ip_prefix}"
                """, language="bash")
                st.success("GÃ¼venlik grubu kuralÄ± ekleme iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "WAF YÃ¶netimi":
    st.header("ğŸ›¡ï¸ WAF YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["WAF OluÅŸtur", "WAF Listele", "WAF Sil", "Kural Ekle"]
    )
    
    if operation == "WAF OluÅŸtur":
        with st.form("waf_create"):
            name = st.text_input("WAF AdÄ±")
            domain = st.text_input("Domain")
            server = st.text_input("Sunucu IP")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="waf_name={name}" -var="domain={domain}" -var="server={server}"
                """, language="bash")
                st.success("WAF oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "WAF Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_waf_instance
                """, language="bash")
            st.success("WAF'lar listeleniyor!")
            
    elif operation == "WAF Sil":
        with st.form("waf_delete"):
            waf_id = st.text_input("WAF ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_waf_instance.{waf_id}
                """, language="bash")
                st.success("WAF silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Kural Ekle":
        with st.form("waf_rule_add"):
            waf_id = st.text_input("WAF ID")
            rule_type = st.selectbox("Kural Tipi", ["blacklist", "whitelist"])
            url = st.text_input("URL")
            action = st.selectbox("Aksiyon", ["block", "allow"])
            
            if st.form_submit_button("Kural Ekle"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="waf_id={waf_id}" -var="rule_type={rule_type}" -var="url={url}" -var="action={action}"
                """, language="bash")
                st.success("WAF kuralÄ± ekleme iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "SSL SertifikalarÄ±":
    st.header("ğŸ”’ SSL SertifikalarÄ± YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["Sertifika OluÅŸtur", "SertifikalarÄ± Listele", "Sertifika Sil", "Sertifika Yenile"]
    )
    
    if operation == "Sertifika OluÅŸtur":
        with st.form("ssl_create"):
            name = st.text_input("Sertifika AdÄ±")
            domain = st.text_input("Domain")
            email = st.text_input("Email")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="ssl_name={name}" -var="domain={domain}" -var="email={email}"
                """, language="bash")
                st.success("SSL sertifikasÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "SertifikalarÄ± Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_scm_certificate
                """, language="bash")
            st.success("SSL sertifikalarÄ± listeleniyor!")
            
    elif operation == "Sertifika Sil":
        with st.form("ssl_delete"):
            cert_id = st.text_input("Sertifika ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_scm_certificate.{cert_id}
                """, language="bash")
                st.success("SSL sertifikasÄ± silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Sertifika Yenile":
        with st.form("ssl_renew"):
            cert_id = st.text_input("Sertifika ID")
            
            if st.form_submit_button("Yenile"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="cert_id={cert_id}" -var="renew=true"
                """, language="bash")
                st.success("SSL sertifikasÄ± yenileme iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "IAM YÃ¶netimi":
    st.header("ğŸ‘¤ IAM YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["KullanÄ±cÄ± OluÅŸtur", "KullanÄ±cÄ±larÄ± Listele", "KullanÄ±cÄ± Sil", "Grup OluÅŸtur", "Rol OluÅŸtur"]
    )
    
    if operation == "KullanÄ±cÄ± OluÅŸtur":
        with st.form("iam_user_create"):
            name = st.text_input("KullanÄ±cÄ± AdÄ±")
            email = st.text_input("Email")
            password = st.text_input("Åifre", type="password")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="user_name={name}" -var="email={email}" -var="password={password}"
                """, language="bash")
                st.success("IAM kullanÄ±cÄ±sÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "KullanÄ±cÄ±larÄ± Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_iam_user
                """, language="bash")
            st.success("IAM kullanÄ±cÄ±larÄ± listeleniyor!")
            
    elif operation == "KullanÄ±cÄ± Sil":
        with st.form("iam_user_delete"):
            user_id = st.text_input("KullanÄ±cÄ± ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_iam_user.{user_id}
                """, language="bash")
                st.success("IAM kullanÄ±cÄ±sÄ± silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Grup OluÅŸtur":
        with st.form("iam_group_create"):
            name = st.text_input("Grup AdÄ±")
            description = st.text_area("AÃ§Ä±klama")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="group_name={name}" -var="description={description}"
                """, language="bash")
                st.success("IAM grubu oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Rol OluÅŸtur":
        with st.form("iam_role_create"):
            name = st.text_input("Rol AdÄ±")
            description = st.text_area("AÃ§Ä±klama")
            policy = st.text_area("Politika (JSON)")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="role_name={name}" -var="description={description}" -var="policy={policy}"
                """, language="bash")
                st.success("IAM rolÃ¼ oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "Anti-DDoS":
    st.header("ğŸ›¡ï¸ Anti-DDoS YÃ¶netimi")
    
    antiddos_action = st.selectbox(
        "Anti-DDoS Ä°ÅŸlemi SeÃ§in",
        ["Koruma OluÅŸtur", "Koruma Listele", "Koruma Sil", "Kural Ekle"]
    )
    
    if antiddos_action == "Koruma OluÅŸtur":
        with st.form("antiddos_form"):
            col1, col2 = st.columns(2)
            with col1:
                protection_name = st.text_input("Koruma AdÄ±")
                ip_address = st.text_input("IP Adresi")
            with col2:
                protection_type = st.selectbox("Koruma Tipi", ["Basic", "Advanced", "Enterprise"])
                bandwidth = st.number_input("Bant GeniÅŸliÄŸi (Mbps)", min_value=100, value=1000)
            
            submitted = st.form_submit_button("Koruma OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="protection_name={protection_name}" -var="ip_address={ip_address}"
                -var="protection_type={protection_type}" -var="bandwidth={bandwidth}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("Anti-DDoS koruma oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Konteyner ve Orkestrasyon Kategorisi
elif service == "ServiceStage":
    st.header("ğŸš€ ServiceStage YÃ¶netimi")
    
    servicestage_action = st.selectbox(
        "ServiceStage Ä°ÅŸlemi SeÃ§in",
        ["Uygulama OluÅŸtur", "Uygulama Listele", "Uygulama Sil", "DaÄŸÄ±tÄ±m Yap"]
    )
    
    if servicestage_action == "Uygulama OluÅŸtur":
        with st.form("servicestage_form"):
            col1, col2 = st.columns(2)
            with col1:
                app_name = st.text_input("Uygulama AdÄ±")
                app_type = st.selectbox("Uygulama Tipi", ["Web", "Microservice", "Function"])
            with col2:
                runtime = st.selectbox("Runtime", ["Java", "Node.js", "Python", "Go"])
                framework = st.selectbox("Framework", ["Spring Boot", "Express", "Flask", "Gin"])
            
            submitted = st.form_submit_button("Uygulama OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="app_name={app_name}" -var="app_type={app_type}"
                -var="runtime={runtime}" -var="framework={framework}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("ServiceStage uygulama oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "SWR YÃ¶netimi":
    st.header("ğŸ³ Software Repository YÃ¶netimi")
    
    swr_action = st.selectbox(
        "SWR Ä°ÅŸlemi SeÃ§in",
        ["Repository OluÅŸtur", "Repository Listele", "Repository Sil", "Image Push"]
    )
    
    if swr_action == "Repository OluÅŸtur":
        with st.form("swr_form"):
            col1, col2 = st.columns(2)
            with col1:
                repo_name = st.text_input("Repository AdÄ±")
                repo_type = st.selectbox("Repository Tipi", ["Private", "Public"])
            with col2:
                description = st.text_area("AÃ§Ä±klama")
                category = st.selectbox("Kategori", ["Application", "Base", "Database"])
            
            submitted = st.form_submit_button("Repository OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="repo_name={repo_name}" -var="repo_type={repo_type}"
                -var="description={description}" -var="category={category}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("SWR repository oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "FunctionGraph":
    st.header("âš¡ FunctionGraph YÃ¶netimi")
    
    function_action = st.selectbox(
        "Function Ä°ÅŸlemi SeÃ§in",
        ["Fonksiyon OluÅŸtur", "Fonksiyon Listele", "Fonksiyon Sil", "Trigger Ekle"]
    )
    
    if function_action == "Fonksiyon OluÅŸtur":
        with st.form("function_form"):
            col1, col2 = st.columns(2)
            with col1:
                function_name = st.text_input("Fonksiyon AdÄ±")
                runtime = st.selectbox("Runtime", ["Python 3.6", "Node.js 12", "Java 8"])
            with col2:
                memory = st.selectbox("Bellek (MB)", ["128", "256", "512", "1024"])
                timeout = st.number_input("Timeout (saniye)", min_value=1, value=30)
            
            submitted = st.form_submit_button("Fonksiyon OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="function_name={function_name}" -var="runtime={runtime}"
                -var="memory={memory}" -var="timeout={timeout}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("FunctionGraph fonksiyonu oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# AI ve BÃ¼yÃ¼k Veri Kategorisi
elif service == "ModelArts":
    st.header("ğŸ¤– ModelArts YÃ¶netimi")
    
    modelarts_action = st.selectbox(
        "ModelArts Ä°ÅŸlemi SeÃ§in",
        ["Notebook OluÅŸtur", "Model EÄŸitimi BaÅŸlat", "Model DaÄŸÄ±tÄ±mÄ± Yap", "Dataset YÃ¼kle"]
    )
    
    if modelarts_action == "Notebook OluÅŸtur":
        with st.form("modelarts_form"):
            col1, col2 = st.columns(2)
            with col1:
                notebook_name = st.text_input("Notebook AdÄ±")
                instance_type = st.selectbox("Instance Tipi", ["CPU", "GPU", "Ascend"])
            with col2:
                framework = st.selectbox("Framework", ["TensorFlow", "PyTorch", "MindSpore"])
                storage = st.number_input("Depolama (GB)", min_value=5, value=20)
            
            submitted = st.form_submit_button("Notebook OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="notebook_name={notebook_name}" -var="instance_type={instance_type}"
                -var="framework={framework}" -var="storage={storage}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("ModelArts notebook oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "DLS YÃ¶netimi":
    st.header("ğŸ“Š Data Lake Service YÃ¶netimi")
    
    dls_action = st.selectbox(
        "DLS Ä°ÅŸlemi SeÃ§in",
        ["Lake OluÅŸtur", "Lake Listele", "Lake Sil", "Dataset Ekle"]
    )
    
    if dls_action == "Lake OluÅŸtur":
        with st.form("dls_form"):
            col1, col2 = st.columns(2)
            with col1:
                lake_name = st.text_input("Lake AdÄ±")
                storage_type = st.selectbox("Depolama Tipi", ["Standard", "Performance"])
            with col2:
                storage_size = st.number_input("Depolama Boyutu (TB)", min_value=1, value=10)
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Lake OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="lake_name={lake_name}" -var="storage_type={storage_type}"
                -var="storage_size={storage_size}" -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DLS lake oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# AI ve BÃ¼yÃ¼k Veri Kategorisi - Eksik Servisler
elif service == "DGC YÃ¶netimi":
    st.header("ğŸ”„ Data Governance Center YÃ¶netimi")
    
    dgc_action = st.selectbox(
        "DGC Ä°ÅŸlemi SeÃ§in",
        ["Workspace OluÅŸtur", "Workspace Listele", "Workspace Sil", "Dataset Ekle"]
    )
    
    if dgc_action == "Workspace OluÅŸtur":
        with st.form("dgc_form"):
            col1, col2 = st.columns(2)
            with col1:
                workspace_name = st.text_input("Workspace AdÄ±")
                workspace_type = st.selectbox("Workspace Tipi", ["Basic", "Professional"])
            with col2:
                description = st.text_area("AÃ§Ä±klama")
                storage_type = st.selectbox("Depolama Tipi", ["Standard", "Performance"])
            
            submitted = st.form_submit_button("Workspace OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="workspace_name={workspace_name}" -var="workspace_type={workspace_type}"
                -var="description={description}" -var="storage_type={storage_type}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DGC workspace oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "CloudTable":
    st.header("ğŸ“Š CloudTable YÃ¶netimi")
    
    cloudtable_action = st.selectbox(
        "CloudTable Ä°ÅŸlemi SeÃ§in",
        ["Cluster OluÅŸtur", "Cluster Listele", "Cluster Sil", "Table OluÅŸtur"]
    )
    
    if cloudtable_action == "Cluster OluÅŸtur":
        with st.form("cloudtable_form"):
            col1, col2 = st.columns(2)
            with col1:
                cluster_name = st.text_input("Cluster AdÄ±")
                cluster_type = st.selectbox("Cluster Tipi", ["HBase", "OpenTSDB"])
            with col2:
                node_count = st.number_input("Node SayÄ±sÄ±", min_value=3, value=3)
                storage_type = st.selectbox("Depolama Tipi", ["SSD", "SAS"])
            
            submitted = st.form_submit_button("Cluster OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="cluster_name={cluster_name}" -var="cluster_type={cluster_type}"
                -var="node_count={node_count}" -var="storage_type={storage_type}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("CloudTable cluster oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "DIS YÃ¶netimi":
    st.header("ğŸ“¥ Data Ingestion Service YÃ¶netimi")
    
    dis_action = st.selectbox(
        "DIS Ä°ÅŸlemi SeÃ§in",
        ["Stream OluÅŸtur", "Stream Listele", "Stream Sil", "Partition Ekle"]
    )
    
    if dis_action == "Stream OluÅŸtur":
        with st.form("dis_form"):
            col1, col2 = st.columns(2)
            with col1:
                stream_name = st.text_input("Stream AdÄ±")
                stream_type = st.selectbox("Stream Tipi", ["COMMON", "ADVANCED"])
            with col2:
                partition_count = st.number_input("Partition SayÄ±sÄ±", min_value=1, value=1)
                data_type = st.selectbox("Veri Tipi", ["BLOB", "JSON", "CSV"])
            
            submitted = st.form_submit_button("Stream OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="stream_name={stream_name}" -var="stream_type={stream_type}"
                -var="partition_count={partition_count}" -var="data_type={data_type}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DIS stream oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# GeliÅŸtirici AraÃ§larÄ± Kategorisi
elif service == "DevCloud":
    st.header("ğŸ› ï¸ DevCloud YÃ¶netimi")
    
    devcloud_action = st.selectbox(
        "DevCloud Ä°ÅŸlemi SeÃ§in",
        ["Proje OluÅŸtur", "Proje Listele", "Proje Sil", "Pipeline OluÅŸtur"]
    )
    
    if devcloud_action == "Proje OluÅŸtur":
        with st.form("devcloud_form"):
            col1, col2 = st.columns(2)
            with col1:
                project_name = st.text_input("Proje AdÄ±")
                project_type = st.selectbox("Proje Tipi", ["Scrum", "Kanban"])
            with col2:
                visibility = st.selectbox("GÃ¶rÃ¼nÃ¼rlÃ¼k", ["Private", "Internal", "Public"])
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Proje OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="project_name={project_name}" -var="project_type={project_type}"
                -var="visibility={visibility}" -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DevCloud proje oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "CodeArts":
    st.header("ğŸ‘¨â€ğŸ’» CodeArts YÃ¶netimi")
    
    codearts_action = st.selectbox(
        "CodeArts Ä°ÅŸlemi SeÃ§in",
        ["Proje OluÅŸtur", "Repository OluÅŸtur", "Pipeline OluÅŸtur", "Build Job OluÅŸtur"]
    )
    
    if codearts_action == "Proje OluÅŸtur":
        with st.form("codearts_form"):
            col1, col2 = st.columns(2)
            with col1:
                project_name = st.text_input("Proje AdÄ±")
                project_type = st.selectbox("Proje Tipi", ["Scrum", "Kanban", "Basic"])
            with col2:
                visibility = st.selectbox("GÃ¶rÃ¼nÃ¼rlÃ¼k", ["Private", "Internal", "Public"])
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Proje OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="project_name={project_name}" -var="project_type={project_type}"
                -var="visibility={visibility}" -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("CodeArts proje oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# GeliÅŸtirici AraÃ§larÄ± Kategorisi - Eksik Servisler
elif service == "ServiceComb":
    st.header("ğŸ”„ ServiceComb YÃ¶netimi")
    
    servicecomb_action = st.selectbox(
        "ServiceComb Ä°ÅŸlemi SeÃ§in",
        ["Mikroservis OluÅŸtur", "Mikroservis Listele", "Mikroservis Sil", "API Ekle"]
    )
    
    if servicecomb_action == "Mikroservis OluÅŸtur":
        with st.form("servicecomb_form"):
            col1, col2 = st.columns(2)
            with col1:
                service_name = st.text_input("Servis AdÄ±")
                framework = st.selectbox("Framework", ["Spring Boot", "Go Chassis", "Java Chassis"])
            with col2:
                version = st.text_input("Versiyon", value="1.0.0")
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Mikroservis OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="service_name={service_name}" -var="framework={framework}"
                -var="version={version}" -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("ServiceComb mikroservis oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "CloudIDE":
    st.header("ğŸ’» CloudIDE YÃ¶netimi")
    
    cloudide_action = st.selectbox(
        "CloudIDE Ä°ÅŸlemi SeÃ§in",
        ["Workspace OluÅŸtur", "Workspace Listele", "Workspace Sil", "Proje Ekle"]
    )
    
    if cloudide_action == "Workspace OluÅŸtur":
        with st.form("cloudide_form"):
            col1, col2 = st.columns(2)
            with col1:
                workspace_name = st.text_input("Workspace AdÄ±")
                template = st.selectbox("Åablon", ["Python", "Java", "Node.js", "Go"])
            with col2:
                instance_type = st.selectbox("Instance Tipi", ["2U4G", "4U8G", "8U16G"])
                storage = st.number_input("Depolama (GB)", min_value=10, value=20)
            
            submitted = st.form_submit_button("Workspace OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="workspace_name={workspace_name}" -var="template={template}"
                -var="instance_type={instance_type}" -var="storage={storage}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("CloudIDE workspace oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Ä°zleme ve YÃ¶netim Kategorisi
elif service == "Cloud Eye":
    st.header("ğŸ‘ï¸ Cloud Eye YÃ¶netimi")
    
    cloudeye_action = st.selectbox(
        "Cloud Eye Ä°ÅŸlemi SeÃ§in",
        ["Alarm KuralÄ± OluÅŸtur", "Dashboard OluÅŸtur", "Metrik GÃ¶rÃ¼ntÃ¼le", "Rapor OluÅŸtur"]
    )
    
    if cloudeye_action == "Alarm KuralÄ± OluÅŸtur":
        with st.form("cloudeye_form"):
            col1, col2 = st.columns(2)
            with col1:
                alarm_name = st.text_input("Alarm AdÄ±")
                metric_name = st.selectbox("Metrik", ["CPU KullanÄ±mÄ±", "Bellek KullanÄ±mÄ±", "Disk KullanÄ±mÄ±"])
            with col2:
                threshold = st.number_input("EÅŸik DeÄŸeri", min_value=0, max_value=100, value=80)
                comparison = st.selectbox("KarÅŸÄ±laÅŸtÄ±rma", [">", ">=", "<", "<=", "="])
            
            submitted = st.form_submit_button("Alarm KuralÄ± OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="alarm_name={alarm_name}" -var="metric_name={metric_name}"
                -var="threshold={threshold}" -var="comparison={comparison}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("Cloud Eye alarm kuralÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "AOM YÃ¶netimi":
    st.header("ğŸ“ˆ Application Operations Management")
    
    aom_action = st.selectbox(
        "AOM Ä°ÅŸlemi SeÃ§in",
        ["Uygulama Ekle", "Uygulama Listele", "Uygulama Sil", "Metrik GÃ¶rÃ¼ntÃ¼le"]
    )
    
    if aom_action == "Uygulama Ekle":
        with st.form("aom_form"):
            col1, col2 = st.columns(2)
            with col1:
                app_name = st.text_input("Uygulama AdÄ±")
                app_type = st.selectbox("Uygulama Tipi", ["Java", "Node.js", "Python", "Go"])
            with col2:
                environment = st.selectbox("Ortam", ["Development", "Testing", "Production"])
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Uygulama Ekle")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="app_name={app_name}" -var="app_type={app_type}"
                -var="environment={environment}" -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("AOM uygulama ekleme iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "LTS YÃ¶netimi":
    st.header("ğŸ“ Log Tank Service YÃ¶netimi")
    
    lts_action = st.selectbox(
        "LTS Ä°ÅŸlemi SeÃ§in",
        ["Log Grubu OluÅŸtur", "Log Grubu Listele", "Log Grubu Sil", "Log Stream Ekle"]
    )
    
    if lts_action == "Log Grubu OluÅŸtur":
        with st.form("lts_form"):
            col1, col2 = st.columns(2)
            with col1:
                group_name = st.text_input("Log Grubu AdÄ±")
                retention = st.selectbox("Saklama SÃ¼resi", ["7 gÃ¼n", "30 gÃ¼n", "90 gÃ¼n"])
            with col2:
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Log Grubu OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="group_name={group_name}" -var="retention={retention}"
                -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("LTS log grubu oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Ä°zleme ve YÃ¶netim Kategorisi - Eksik Servisler
elif service == "Terraform Durumu":
    st.header("ğŸ”„ Terraform Durumu")
    
    terraform_action = st.selectbox(
        "Terraform Ä°ÅŸlemi SeÃ§in",
        ["Durum GÃ¶rÃ¼ntÃ¼le", "Plan OluÅŸtur", "DeÄŸiÅŸiklikleri Uygula", "KaynaklarÄ± Temizle"]
    )
    
    if terraform_action == "Durum GÃ¶rÃ¼ntÃ¼le":
        with st.form("terraform_form"):
            col1, col2 = st.columns(2)
            with col1:
                show_details = st.checkbox("DetaylarÄ± GÃ¶ster", value=True)
                output_format = st.selectbox("Ã‡Ä±ktÄ± FormatÄ±", ["table", "json", "raw"])
            with col2:
                refresh = st.checkbox("Durumu Yenile", value=False)
            
            submitted = st.form_submit_button("Durum GÃ¶rÃ¼ntÃ¼le")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform state list
                """
                if refresh:
                    terraform_cmd += "\nterraform refresh"
                st.code(terraform_cmd, language="bash")
                st.success("Terraform durumu gÃ¶rÃ¼ntÃ¼leme iÅŸlemi baÅŸlatÄ±ldÄ±!")

# GÃ¼venlik ve Kimlik Kategorisi - Eksik Servisler
elif service == "KMS YÃ¶netimi":
    st.header("ğŸ” Key Management Service YÃ¶netimi")
    
    kms_action = st.selectbox(
        "KMS Ä°ÅŸlemi SeÃ§in",
        ["Anahtar OluÅŸtur", "Anahtar Listele", "Anahtar Sil", "Åifreleme/Åifre Ã‡Ã¶zme"]
    )
    
    if kms_action == "Anahtar OluÅŸtur":
        with st.form("kms_form"):
            col1, col2 = st.columns(2)
            with col1:
                key_name = st.text_input("Anahtar AdÄ±")
                key_type = st.selectbox("Anahtar Tipi", ["AES_256", "RSA_2048", "RSA_4096"])
            with col2:
                key_usage = st.selectbox("KullanÄ±m AmacÄ±", ["ENCRYPT_DECRYPT", "SIGN_VERIFY"])
                description = st.text_area("AÃ§Ä±klama")
            
            submitted = st.form_submit_button("Anahtar OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="key_name={key_name}" -var="key_type={key_type}"
                -var="key_usage={key_usage}" -var="description={description}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("KMS anahtarÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "CBR YÃ¶netimi":
    st.header("ğŸ’¾ Cloud Backup and Recovery YÃ¶netimi")
    
    cbr_action = st.selectbox(
        "CBR Ä°ÅŸlemi SeÃ§in",
        ["Yedekleme PolitikasÄ± OluÅŸtur", "Yedekleme Listele", "Yedekleme Sil", "Geri YÃ¼kleme"]
    )
    
    if cbr_action == "Yedekleme PolitikasÄ± OluÅŸtur":
        with st.form("cbr_form"):
            col1, col2 = st.columns(2)
            with col1:
                policy_name = st.text_input("Politika AdÄ±")
                backup_type = st.selectbox("Yedekleme Tipi", ["VM", "Disk", "SFS"])
            with col2:
                retention_days = st.number_input("Saklama SÃ¼resi (GÃ¼n)", min_value=1, value=7)
                schedule = st.selectbox("Zamanlama", ["GÃ¼nlÃ¼k", "HaftalÄ±k", "AylÄ±k"])
            
            submitted = st.form_submit_button("Politika OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="policy_name={policy_name}" -var="backup_type={backup_type}"
                -var="retention_days={retention_days}" -var="schedule={schedule}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("CBR yedekleme politikasÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Konteyner ve Orkestrasyon Kategorisi - Eksik Servisler
elif service == "CCE YÃ¶netimi":
    st.header("ğŸ³ Cloud Container Engine YÃ¶netimi")
    
    cce_action = st.selectbox(
        "CCE Ä°ÅŸlemi SeÃ§in",
        ["Cluster OluÅŸtur", "Cluster Listele", "Cluster Sil", "Node Pool Ekle"]
    )
    
    if cce_action == "Cluster OluÅŸtur":
        with st.form("cce_form"):
            col1, col2 = st.columns(2)
            with col1:
                cluster_name = st.text_input("Cluster AdÄ±")
                cluster_type = st.selectbox("Cluster Tipi", ["VirtualMachine", "BareMetal"])
            with col2:
                version = st.selectbox("Kubernetes Versiyonu", ["v1.19", "v1.21", "v1.23"])
                node_count = st.number_input("Node SayÄ±sÄ±", min_value=1, value=3)
            
            submitted = st.form_submit_button("Cluster OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="cluster_name={cluster_name}" -var="cluster_type={cluster_type}"
                -var="version={version}" -var="node_count={node_count}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("CCE cluster oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Depolama ve VeritabanÄ± Kategorisi - Eksik Servisler
elif service == "DCS YÃ¶netimi":
    st.header("ğŸ—„ï¸ Distributed Cache Service YÃ¶netimi")
    
    dcs_action = st.selectbox(
        "DCS Ä°ÅŸlemi SeÃ§in",
        ["Instance OluÅŸtur", "Instance Listele", "Instance Sil", "Yedekleme OluÅŸtur"]
    )
    
    if dcs_action == "Instance OluÅŸtur":
        with st.form("dcs_form"):
            col1, col2 = st.columns(2)
            with col1:
                instance_name = st.text_input("Instance AdÄ±")
                engine = st.selectbox("Cache Motoru", ["Redis", "Memcached"])
            with col2:
                capacity = st.selectbox("Kapasite", ["2GB", "4GB", "8GB", "16GB"])
                flavor = st.selectbox("Instance Tipi", ["Single", "Master-Slave", "Cluster"])
            
            submitted = st.form_submit_button("Instance OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="instance_name={instance_name}" -var="engine={engine}"
                -var="capacity={capacity}" -var="flavor={flavor}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DCS instance oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "DMS YÃ¶netimi":
    st.header("ğŸ“¨ Distributed Message Service YÃ¶netimi")
    
    dms_action = st.selectbox(
        "DMS Ä°ÅŸlemi SeÃ§in",
        ["Instance OluÅŸtur", "Instance Listele", "Instance Sil", "Topic OluÅŸtur"]
    )
    
    if dms_action == "Instance OluÅŸtur":
        with st.form("dms_form"):
            col1, col2 = st.columns(2)
            with col1:
                instance_name = st.text_input("Instance AdÄ±")
                engine = st.selectbox("Mesaj Motoru", ["Kafka", "RabbitMQ"])
            with col2:
                version = st.selectbox("Versiyon", ["2.3.0", "2.7.0", "3.7.0"])
                flavor = st.selectbox("Instance Tipi", ["Single", "Cluster"])
            
            submitted = st.form_submit_button("Instance OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="instance_name={instance_name}" -var="engine={engine}"
                -var="version={version}" -var="flavor={flavor}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("DMS instance oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Hesaplama ve AÄŸ Kategorisi - Eksik Servisler
elif service == "ELB YÃ¶netimi":
    st.header("âš–ï¸ Elastic Load Balancer YÃ¶netimi")
    
    elb_action = st.selectbox(
        "ELB Ä°ÅŸlemi SeÃ§in",
        ["Load Balancer OluÅŸtur", "Load Balancer Listele", "Load Balancer Sil", "Listener Ekle"]
    )
    
    if elb_action == "Load Balancer OluÅŸtur":
        with st.form("elb_form"):
            col1, col2 = st.columns(2)
            with col1:
                lb_name = st.text_input("Load Balancer AdÄ±")
                lb_type = st.selectbox("Load Balancer Tipi", ["Application", "Network"])
            with col2:
                bandwidth = st.number_input("Bant GeniÅŸliÄŸi (Mbps)", min_value=1, value=100)
                flavor = st.selectbox("Instance Tipi", ["small", "medium", "large"])
            
            submitted = st.form_submit_button("Load Balancer OluÅŸtur")
            if submitted:
                terraform_cmd = f"""
                terraform init
                terraform apply -auto-approve -var="lb_name={lb_name}" -var="lb_type={lb_type}"
                -var="bandwidth={bandwidth}" -var="flavor={flavor}"
                """
                st.code(terraform_cmd, language="bash")
                st.success("ELB load balancer oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")

# AI ve BÃ¼yÃ¼k Veri Kategorisi - Eksik Servisler
elif service == "MRS YÃ¶netimi":
    st.header("ğŸ“Š MapReduce Service YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["Cluster OluÅŸtur", "Cluster Listele", "Cluster Sil", "Job GÃ¶nder"]
    )
    
    if operation == "Cluster OluÅŸtur":
        with st.form("mrs_cluster_create"):
            name = st.text_input("Cluster AdÄ±")
            cluster_type = st.selectbox("Cluster Tipi", ["analysis", "streaming", "hybrid"])
            node_count = st.number_input("Node SayÄ±sÄ±", min_value=1)
            node_flavor = st.text_input("Node Flavor")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="cluster_name={name}" -var="cluster_type={cluster_type}" -var="node_count={node_count}" -var="node_flavor={node_flavor}"
                """, language="bash")
                st.success("MRS cluster oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Cluster Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_mrs_cluster
                """, language="bash")
            st.success("MRS cluster'larÄ± listeleniyor!")
            
    elif operation == "Cluster Sil":
        with st.form("mrs_cluster_delete"):
            cluster_id = st.text_input("Cluster ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_mrs_cluster.{cluster_id}
                """, language="bash")
                st.success("MRS cluster silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Job GÃ¶nder":
        with st.form("mrs_job_submit"):
            cluster_id = st.text_input("Cluster ID")
            job_type = st.selectbox("Job Tipi", ["MapReduce", "Spark", "Hive", "SparkScript"])
            job_name = st.text_input("Job AdÄ±")
            jar_path = st.text_input("JAR Dosya Yolu")
            
            if st.form_submit_button("Job GÃ¶nder"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="cluster_id={cluster_id}" -var="job_type={job_type}" -var="job_name={job_name}" -var="jar_path={jar_path}"
                """, language="bash")
                st.success("MRS job gÃ¶nderme iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "API Gateway":
    st.header("ğŸŒ API Gateway YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["API OluÅŸtur", "API Listele", "API Sil", "Environment Ekle"]
    )
    
    if operation == "API OluÅŸtur":
        with st.form("api_create"):
            name = st.text_input("API AdÄ±")
            group_id = st.text_input("API Group ID")
            path = st.text_input("Path")
            method = st.selectbox("Method", ["GET", "POST", "PUT", "DELETE"])
            backend_type = st.selectbox("Backend Tipi", ["HTTP", "FUNCTION"])
            backend_url = st.text_input("Backend URL")
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="api_name={name}" -var="group_id={group_id}" -var="path={path}" -var="method={method}" -var="backend_type={backend_type}" -var="backend_url={backend_url}"
                """, language="bash")
                st.success("API oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "API Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_apig_api
                """, language="bash")
            st.success("API'ler listeleniyor!")
            
    elif operation == "API Sil":
        with st.form("api_delete"):
            api_id = st.text_input("API ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_apig_api.{api_id}
                """, language="bash")
                st.success("API silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Environment Ekle":
        with st.form("api_env_add"):
            group_id = st.text_input("API Group ID")
            env_name = st.text_input("Environment AdÄ±")
            env_description = st.text_area("Environment AÃ§Ä±klamasÄ±")
            
            if st.form_submit_button("Environment Ekle"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="group_id={group_id}" -var="env_name={env_name}" -var="env_description={env_description}"
                """, language="bash")
                st.success("API environment ekleme iÅŸlemi baÅŸlatÄ±ldÄ±!")

elif service == "CES YÃ¶netimi":
    st.header("ğŸ“Š Cloud Eye Service YÃ¶netimi")
    
    # Ä°ÅŸlem seÃ§imi
    operation = st.selectbox(
        "Ä°ÅŸlem SeÃ§in",
        ["Alarm KuralÄ± OluÅŸtur", "Alarm KurallarÄ±nÄ± Listele", "Alarm KuralÄ± Sil", "Metrik GÃ¶rÃ¼ntÃ¼le"]
    )
    
    if operation == "Alarm KuralÄ± OluÅŸtur":
        with st.form("ces_alarm_create"):
            name = st.text_input("Alarm KuralÄ± AdÄ±")
            metric_name = st.text_input("Metrik AdÄ±")
            threshold = st.number_input("EÅŸik DeÄŸeri")
            comparison_operator = st.selectbox("KarÅŸÄ±laÅŸtÄ±rma OperatÃ¶rÃ¼", [">", "<", ">=", "<=", "="])
            period = st.number_input("Periyot (saniye)", min_value=1)
            
            if st.form_submit_button("OluÅŸtur"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="alarm_name={name}" -var="metric_name={metric_name}" -var="threshold={threshold}" -var="comparison_operator={comparison_operator}" -var="period={period}"
                """, language="bash")
                st.success("CES alarm kuralÄ± oluÅŸturma iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Alarm KurallarÄ±nÄ± Listele":
        if st.button("Listele"):
            st.info("Terraform komutu oluÅŸturuluyor...")
            st.code("""
terraform init
terraform state list | grep huaweicloud_ces_alarmrule
                """, language="bash")
            st.success("CES alarm kurallarÄ± listeleniyor!")
            
    elif operation == "Alarm KuralÄ± Sil":
        with st.form("ces_alarm_delete"):
            alarm_id = st.text_input("Alarm KuralÄ± ID")
            
            if st.form_submit_button("Sil"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform destroy -target=huaweicloud_ces_alarmrule.{alarm_id}
                """, language="bash")
                st.success("CES alarm kuralÄ± silme iÅŸlemi baÅŸlatÄ±ldÄ±!")
                
    elif operation == "Metrik GÃ¶rÃ¼ntÃ¼le":
        with st.form("ces_metric_view"):
            metric_name = st.text_input("Metrik AdÄ±")
            start_time = st.text_input("BaÅŸlangÄ±Ã§ ZamanÄ± (YYYY-MM-DD HH:MM:SS)")
            end_time = st.text_input("BitiÅŸ ZamanÄ± (YYYY-MM-DD HH:MM:SS)")
            
            if st.form_submit_button("GÃ¶rÃ¼ntÃ¼le"):
                st.info("Terraform komutu oluÅŸturuluyor...")
                st.code(f"""
terraform init
terraform apply -var="metric_name={metric_name}" -var="start_time={start_time}" -var="end_time={end_time}"
                """, language="bash")
                st.success("CES metrik gÃ¶rÃ¼ntÃ¼leme iÅŸlemi baÅŸlatÄ±ldÄ±!")

# Alt bilgi
st.markdown("---")
st.markdown("""
<div style='text-align: center'>
    <p>Huawei Cloud YÃ¶netim Paneli v2.0 | GeliÅŸtirici: Berk Kaya</p>
    <p>Son gÃ¼ncelleme: {}</p>
</div>
""".format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")), unsafe_allow_html=True) 